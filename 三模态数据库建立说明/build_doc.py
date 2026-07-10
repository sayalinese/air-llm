# -*- coding: utf-8 -*-
"""Generate Aeolus_V2 documentation Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# Page setup
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Helper functions
def add_table_row(table, cells_text, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        p = row.cells[i].paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        if header:
            run.bold = True
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D5E8F0')
            shading.set(qn('w:val'), 'clear')
            row.cells[i]._tc.get_or_add_tcPr().append(shading)
        if bold:
            run.bold = True

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    shading.set(qn('w:val'), 'clear')
    p._p.get_or_add_pPr().append(shading)

# ===== Cover Page =====
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Aeolus_V2')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('航班延误预测三模态数据库建立说明')
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ['数据源: BTS On-Time Performance + Meteostat',
             '时间范围: 2016-2025',
             '机场数: 322个',
             '',
             '2026年5月']:
    run = info.add_run(line + '\n')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ===== Table of Contents =====
doc.add_heading('目录', level=1)
toc = doc.add_paragraph()
run = toc.add_run('（在Word中右键此处 → 更新域 → 更新整个目录）')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# Create TOC field
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')

run2 = doc.add_paragraph().add_run()
run2._r.append(fldChar1)
run2._r.append(instrText)
run2._r.append(fldChar2)
run2._r.append(fldChar3)

doc.add_page_break()

# ===== 1. Overview =====
doc.add_heading('1. 概述', level=1)
doc.add_paragraph(
    'Aeolus_V2 是对原版 Aeolus 航班延误预测数据集的全面升级版。本数据集覆盖 2016-2025 年'
    '美国国内航班，包含三种模态:'
)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Flight_Tabular: ').bold = True
p.add_run('结构化表格数据，对齐原版特征工程并扩展，每文件29列，含机场等级与同机前序特征')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Flight_Chain: ').bold = True
p.add_run('按 TAIL_NUM 分组的航班序列，追踪同一物理飞机的一日航迹')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Flight_Network: ').bold = True
p.add_run('机场级拓扑图，DFS树结构，捕捉延误溢出传播')

doc.add_paragraph()
doc.add_paragraph('V2 核心改进:')

items = [
    '链分组键从 (OP_CARRIER, FL_NUM, DATE) 改为 (TAIL_NUM, DATE)，正确追踪同一物理飞机',
    '网络边特征中 AIRCRAFT_NUM 和 AIRPORT 正确赋值（原版均为 1 的 bug 已修复）',
    '三种模态均保留 TAIL_NUM 作为 key',
    '全局编码（按年统一编码），确保跨天一致性',
    '时间范围扩展到 2016-2025（原版至 2024）',
    'Tabular 新增机场等级特征 (ORIGIN_TIER/DEST_TIER)，来自 OurAirports 机场类型字段',
    'Tabular 新增前序航班特征 (PREV_DEP_DELAY/PREV_ARR_DELAY/PREV_TIME_GAP)，向量化 shift() 实现',
    '统一使用 DEP_DELAY 作为目标标签，三模态一致',
    'Network 改用迭代栈 DFS 避免递归栈溢出',
    '每年度输出一份 _tails.json 尾号映射表，便于反查原始飞机编号',
]
for i, item in enumerate(items, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(item)

doc.add_page_break()

# ===== 2. Data Sources =====
doc.add_heading('2. 数据源说明', level=1)

doc.add_heading('2.1 BTS On-Time Performance', level=2)
p = doc.add_paragraph()
p.add_run('下载地址: ').bold = True
p.add_run('https://transtats.bts.gov/PREZIP/')
doc.add_paragraph('内容: 美国国内航班准点报告，包含出发/到达时间、延误、取消等字段')
doc.add_paragraph('格式: 月度ZIP压缩包，内含CSV文件，列名为驼峰命名（如 Tail_Number 对应 TAIL_NUM）')
p = doc.add_paragraph()
p.add_run('下载文件: ').bold = True
p.add_run('120 个 ZIP 文件（每月一个，2016-01 ~ 2025-12），总计约 2.5 GB')

doc.add_heading('2.2 Meteostat 天气数据', level=2)
p = doc.add_paragraph()
p.add_run('库: ').bold = True
p.add_run('meteostat Python 包')
doc.add_paragraph('内容: 逐小时气温 (TEMP)、降水量 (PRCP)、风速 (WSPD)')
doc.add_paragraph('覆盖机场: 322 个（基于 BTS 数据中出现的所有机场）')

doc.add_heading('2.3 机场坐标', level=2)
p = doc.add_paragraph()
p.add_run('数据源: ').bold = True
p.add_run('OurAirports 开源数据库 (CC-BY 4.0)')
doc.add_paragraph('用途: 匹配 BTS 航班的 ORIGIN/DEST 机场经纬度')

doc.add_page_break()

# ===== 3. Pipeline Overview =====
doc.add_heading('3. 整体流程', level=1)
doc.add_paragraph('下表展示从原始数据到三模态数据集的完整流程:')

# Pipeline table
t = doc.add_table(rows=10, cols=2)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['步骤', '说明']
data = [
    ['1. BTS ZIP下载', 'download_bts_fixed.py → 120个ZIP (bts_2016_2025/)'],
    ['2. ZIP→按日CSV', 'process_bts_zips.py → 3653个日CSV (raw/bts/)'],
    ['3. 天气数据下载', 'fetch_weather_mt.py → 322机场天气 (raw/weather/)'],
    ['4. 合并天气+坐标', 'merge_flight_weather.py → flight_with_weather/'],
    ['5. Flight_Tabular', 'build_tabular_v2.py → 特征工程 → 按日输出'],
    ['6. Flight_Chain', 'build_chain_fast.py → 全年编码 → 按日.pt输出'],
    ['7. Flight_Network', 'build_network_mt.py → 全局编码 → 按日.dgl输出'],
]
for j, h in enumerate(headers):
    t.rows[0].cells[j].text = h
    for p in t.rows[0].cells[j].paragraphs:
        for r in p.runs:
            r.bold = True
for i, row_data in enumerate(data, 1):
    for j, cell_text in enumerate(row_data):
        t.rows[i].cells[j].text = cell_text

doc.add_paragraph()
doc.add_heading('目录结构', level=3)
add_code('bts_2016_2025/                   # 120 个 BTS ZIP 原始文件')
add_code('Aeolus_V2/raw/bts/               # 按日CSV (2016/01/2016-01-01.csv)')
add_code('Aeolus_V2/raw/weather/           # 322 个机场天气缓存')
add_code('flight_with_weather/             # 融合后的日CSV')
add_code('Aeolus_V2/dataset/')
add_code('  Flight_Tabular/   # 特征工程后的日CSV (29列，含机场等级与前序特征)')
add_code('  Flight_Chain/     # 每日 .pt 文件')
add_code('  Flight_Network/   # 每日 .dgl 图文件')

doc.add_page_break()

# ===== 4. Step by Step =====
doc.add_heading('4. 步骤详解', level=1)

# Step 1
doc.add_heading('4.1 下载 BTS ZIP', level=2)
p = doc.add_paragraph()
p.add_run('脚本: ').bold = True
p.add_run('download_bts_fixed.py').font.name = 'Consolas'
doc.add_paragraph('关键设计:')
items = [
    'verify=False 跳过SSL证书验证（解决 SSLEOFError）',
    '已下载文件自动跳过（断点续传）',
    '1MB分块流式下载，防止连接断开',
    '最多 5 次重试，等待时间递增',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
p = doc.add_paragraph()
p.add_run('运行: ').bold = True
p.add_run('python download_bts_fixed.py').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('产出: ').bold = True
p.add_run('bts_2016_2025/ 下 120 个 ZIP 文件（约 2.5 GB）')

# Step 2
doc.add_heading('4.2 ZIP → 按日 CSV', level=2)
p = doc.add_paragraph()
p.add_run('脚本: ').bold = True
p.add_run('process_bts_zips.py').font.name = 'Consolas'
doc.add_paragraph('关键设计:')
items = [
    '解压ZIP → 匹配BTS新版驼峰列名 → 过滤保留字段',
    '暂存到 BytesIO 解决 seek on closed file 问题',
    '按天分割保存为单个 CSV，输出到 raw/bts/{year}/{month}/',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
p = doc.add_paragraph()
p.add_run('产出: ').bold = True
p.add_run('3,653 个日CSV文件（约 5600 万条记录）')

# Step 3
doc.add_heading('4.3 下载天气数据', level=2)
p = doc.add_paragraph()
p.add_run('脚本: ').bold = True
p.add_run('fetch_weather_mt.py').font.name = 'Consolas'
doc.add_paragraph('关键设计:')
items = [
    '从已有 BTS CSV 扫描所有机场代码',
    '使用 Meteostat 库查询每个机场的逐小时天气 (2016-2025)',
    '4 线程并行下载（缓存为 wx_{ICAO}.parquet）',
    '如果缺少机场，可运行 find_airports_now.py 补下',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
p = doc.add_paragraph()
p.add_run('产出: ').bold = True
p.add_run('raw/weather/ 下 322 个 wx_*.parquet 文件（约 2800 万条记录）')

# Step 4
doc.add_heading('4.4 合并航班 + 天气 + 机场坐标', level=2)
p = doc.add_paragraph()
p.add_run('脚本: ').bold = True
p.add_run('merge_flight_weather.py').font.name = 'Consolas'
doc.add_paragraph('关键设计:')
items = [
    '按年处理（一年只做2次merge，避免逐文件O(N²)开销）',
    '数据格式化：日期转字符串、小时提取为整数',
    '用 map() 替代 merge() 避免 dtype 冲突',
    '缺失天气值用 0 填充',
    '按日写出到 flight_with_weather/{year}/{month}/',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
p = doc.add_paragraph()
p.add_run('产出: ').bold = True
p.add_run('flight_with_weather/ 下 3,653 个日CSV，每文件34列')

doc.add_page_break()

# Step 5
doc.add_heading('4.5 Flight_Tabular — 特征工程', level=2)
p = doc.add_paragraph()
p.add_run('脚本: ').bold = True
p.add_run('build_tabular_v2.py').font.name = 'Consolas'
doc.add_paragraph('对齐原版 Flight_tab.py 并增加 V2 改进特征:')
items = [
    '离群值剔除: DEP_DELAY 1%-99% 百分位（目标统一为起飞延误）',
    '日期处理: FL_DATE → FL_YEAR, MONTH→FL_MONTH, DAY_OF_WEEK→FL_WEEK',
    '时间转换: CRS_DEP_TIME(HHMM) → CRS_DEP_TIME_MIN(分钟数)',
    'ORIGIN_INDEX = ORIGIN, DEST_INDEX = DEST（保持字符串，不做编码）',
    '不做归一化（与原版一致）',
    '机场等级特征: 从 OurAirports 数据库的 type 字段映射，large_airport=3(枢纽), '
    'medium_airport=2, small_airport/其他=1，增加 ORIGIN_TIER/DEST_TIER 两列',
    '前序航班特征: 按 (TAIL_NUM, DATE) 分组，用 pandas shift() 向量化提取同机前序航班的 '
    'PREV_DEP_DELAY(前序起飞延误)、PREV_ARR_DELAY(前序到达延误)、PREV_TIME_GAP(间隔分钟)，'
    '无前序时填充 0',
    '目标列统一: 仅保留 DEP_DELAY 作为目标，移除 ARR_DELAY，与 Chain/Network 保持一致',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

p = doc.add_paragraph()
p.add_run('输出: ').bold = True
p.add_run('Flight_Tabular/{year}/{month}/flight_with_weather_{YYMMDD}.csv')
doc.add_paragraph('每文件 29 列: Target(1: DEP_DELAY) + Categorical(8) + Tier(2) + Continuous(14) + Prev(3) + TAIL_NUM')
doc.add_paragraph('列顺序: DEP_DELAY, OP_CARRIER, OP_CARRIER_FL_NUM, FL_YEAR, FL_MONTH, FL_DAY, '
    'FL_WEEK, ORIGIN_INDEX, DEST_INDEX, ORIGIN_TIER, DEST_TIER, '
    'CRS_DEP_TIME_MIN, CRS_ARR_TIME_MIN, CRS_ELAPSED_TIME, FLIGHTS, '
    'PREV_DEP_DELAY, PREV_ARR_DELAY, PREV_TIME_GAP, '
    'O_TEMP, O_PRCP, O_WSPD, D_TEMP, D_PRCP, D_WSPD, '
    'O_LATITUDE, O_LONGITUDE, D_LATITUDE, D_LONGITUDE, TAIL_NUM')

# Step 6
doc.add_heading('4.6 Flight_Chain — 同机航班序列', level=2)
p = doc.add_paragraph()
p.add_run('脚本: ').bold = True
p.add_run('build_chain_fast.py').font.name = 'Consolas'
doc.add_paragraph('关键设计:')
items = [
    '按年处理: 加载当年所有日CSV → 统一编码 → 逐日写出',
    '全局编码: 全年 LabelEncoder 统一编码 OP_CARRIER、航班号、机场、TAIL_NUM',
    '分组键: (TAIL_NUM, DATE) — V2改进，追踪同一物理飞机当天航班序列',
    '最大序列长度: 6（padding到统一长度）',
    '每条 chain: dense_feat(7维) + sparse_feat(9维,含TAIL_NUM_ENC) + labels(1维:DEP>15) + delays(1维:DEP_DELAY) + valid_len',
    '存储格式: 每日一个 .pt 文件（dict格式）',
    '尾号映射: 每年度输出 {year}_tails.json，可用于将 TAIL_NUM_ENC 反查为原始尾号',
]
for item in items:
    doc.add_paragraph(item, style='List Number')
p = doc.add_paragraph()
p.add_run('输出: ').bold = True
p.add_run('Flight_Chain/{year}/{month}/flight_chain_{YYMMDD}.pt（附 {year}_tails.json）')

# Step 7
doc.add_heading('4.7 Flight_Network — 延误传播图', level=2)
p = doc.add_paragraph()
p.add_run('脚本: ').bold = True
p.add_run('build_network_mt.py').font.name = 'Consolas'
doc.add_paragraph('关键设计:')
items = [
    '按年处理: 先扫描全年数据建全局机场映射表和尾号映射表',
    'CRS_DEP_TIME/CRS_ARR_TIME 保持 HHMM 整数（不做 datetime 转换）',
    '建图算法: 迭代栈 DFS（替代原版递归 DFS），避免 Python 递归栈溢出',
    '15 分钟转机窗口: 同一机场降落→起飞之间 15 分钟内（HHMM 整数进位处理）',
    '节点特征: 25 个 ndata（22独立 + feat 15维 + label + TAIL_NUM_ENC）',
    'label = DEP_DELAY/60（统一使用起飞延误作为标签）',
    '边特征: INTERVAL_TIME(float32), AIRCRAFT_NUM(int16), AIRPORT(int16)',
    'V2 修复: 边特征 AIRCRAFT_NUM = 源飞机 TAIL_NUM_ENC, AIRPORT = DEST 编码',
    '尾号映射: 每年度输出 {year}_tails.json，便于反查原始尾号',
]
for item in items:
    doc.add_paragraph(item, style='List Number')
p = doc.add_paragraph()
p.add_run('输出: ').bold = True
p.add_run('Flight_Network/{year}/{month}/flight_network_{YYMMDD}.dgl（附 {year}_tails.json）')

doc.add_page_break()

# ===== 5. Modality Comparison =====
doc.add_heading('5. 三模态对比', level=1)
doc.add_paragraph('下表展示 V2 与原版的关键差异:')

t = doc.add_table(rows=11, cols=4)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
table_data = [
    ['维度', 'Flight_Tabular', 'Flight_Chain', 'Flight_Network'],
    ['分组键', 'N/A', '(TAIL_NUM, DATE) V2', 'N/A'],
    ['数据源', 'flight_with_weather/', 'flight_with_weather/', 'flight_with_weather/'],
    ['编码方式', '保持原始字符串', '全年 LabelEncoder', '全年建映射表'],
    ['异常值', 'DEP_DELAY 1%-99% 剔除', '不剔除', '不剔除'],
    ['每文件字段', '29列(含Tier+Prev)', '7d+9s+1label+1delay', '25 ndata'],
    ['TAIL_NUM', '保留字符串', 'sparse_feat第9维', "ndata['TAIL_NUM_ENC']"],
    ['标签统一', 'DEP_DELAY', '(DEP>15) 单维', 'DEP_DELAY/60'],
    ['尾号映射表', 'N/A（字符串直接可读）', '{year}_tails.json', '{year}_tails.json'],
    ['存储格式', 'CSV', '.pt (dict)', '.dgl (DGL图)'],
]
for i, row_data in enumerate(table_data):
    for j, cell_text in enumerate(row_data):
        t.rows[i].cells[j].text = cell_text
        if i == 0:
            for p in t.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_page_break()

# ===== 6. File List =====
doc.add_heading('6. 附录: 脚本文件清单', level=1)
doc.add_paragraph('以下文件位于 scripts/ 文件夹下:')

t = doc.add_table(rows=11, cols=4)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
file_data = [
    ['文件名', '步骤', '功能', '输入→输出'],
    ['download_bts_fixed.py', '1', '下载BTS', 'BTS网站→ZIP文件'],
    ['verify_zips.py', '1批', '验证ZIP', 'ZIP→完整性报告'],
    ['process_bts_zips.py', '2', '转每日CSV', 'ZIP→每日CSV'],
    ['find_airports_now.py', '3', '扫描机场', 'CSV→机场列表'],
    ['fetch_weather_mt.py', '3', '下载天气', '机场→天气缓存'],
    ['merge_flight_weather.py', '4', '合并数据', 'CSV+天气→融合CSV'],
    ['build_tabular_v2.py', '5', 'Tabular', '融合CSV→特征工程CSV'],
    ['build_chain_fast.py', '6', 'Chain', '融合CSV→.pt'],
    ['build_network_mt.py', '7', 'Network', '融合CSV→.dgl'],
]
for i, row_data in enumerate(file_data):
    for j, cell_text in enumerate(row_data):
        t.rows[i].cells[j].text = cell_text
        if i == 0:
            for p in t.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('注意: 所有脚本使用绝对路径，运行时需确保目录结构与文档描述一致。')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
out_path = r'd:\Daisy\Aeolus_V2\三模态数据库建立说明\Aeolus_V2_三模态数据库建立说明.docx'
doc.save(out_path)
print('Done: ' + out_path)
