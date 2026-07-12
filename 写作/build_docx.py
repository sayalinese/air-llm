#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 写作/ 目录下的 Markdown 章节文件合并为 Word 文档
格式遵循 chinese-journal-writing skill 规范

运行：
    D:\\vllm\\python\\python.exe build_docx.py              # 生成全文
    D:\\vllm\\python\\python.exe build_docx.py 01_绪论.md    # 只生成单章
"""

import re
import sys
from pathlib import Path

# ── 依赖检查 ──────────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    print("正在安装 python-docx ...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── 路径配置 ──────────────────────────────────────────────────────────────────
WRITING_DIR  = Path(__file__).parent
OUTPUT_FILE   = WRITING_DIR / "论文全文.docx"
ABSTRACT_FILE = "00_摘要.md"
CHAPTER_FILES = [
    "01_绪论.md",
    "02_相关理论基础.md",
    "03_多元数据处理.md",
    "04_模型设计.md",
    "05_实验评估.md",
    "06_结论与展望.md",
]

# ── 字体 ─────────────────────────────────────────────────────────────────────
FONT_SONG = "宋体"
FONT_HEI  = "黑体"
FONT_KAI  = "楷体"
FONT_EN   = "Times New Roman"


# ════════════════════════════════════════════════════════════════════════════
#  基础工具
# ════════════════════════════════════════════════════════════════════════════

def _cn_font(run, cn=FONT_SONG, en=FONT_EN):
    """设置中英文字体"""
    rPr = run._r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    rf.set(qn("w:eastAsia"), cn)
    rf.set(qn("w:ascii"), en)
    rf.set(qn("w:hAnsi"), en)


def _fmt_run(run, cn=FONT_SONG, size=12, bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    _cn_font(run, cn)


def _spacing(p, before=0, after=0, lines=1.5):
    pPr = p._p.get_or_add_pPr()
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing"); pPr.append(sp)
    sp.set(qn("w:line"),     str(int(lines * 240)))
    sp.set(qn("w:lineRule"), "auto")
    sp.set(qn("w:before"),   str(int(before * 20)))
    sp.set(qn("w:after"),    str(int(after  * 20)))


def _indent(p, first_cm=0.847, left_cm=0):
    """首行缩进 / 左缩进（单位 cm，内部转 twips）"""
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind"); pPr.append(ind)
    if first_cm:
        ind.set(qn("w:firstLine"), str(int(first_cm * 567)))
    if left_cm:
        ind.set(qn("w:left"), str(int(left_cm * 567)))


def _page_break_before(p):
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pageBreakBefore")
    pb.set(qn("w:val"), "1")
    pPr.append(pb)


# ════════════════════════════════════════════════════════════════════════════
#  内联文本渲染（处理 \( ... \) 内联数学）
# ════════════════════════════════════════════════════════════════════════════

def _render_inline(paragraph, text):
    """渲染内联内容：数学公式 + <sup>上标"""
    # 先按数学公式拆分
    parts = re.split(r"(\\\(.*?\\\))", text, flags=re.DOTALL)
    for part in parts:
        if part.startswith("\\(") and part.endswith("\\)"):
            run = paragraph.add_run(part[2:-2].strip())
            _fmt_run(run, FONT_SONG, size=12, italic=True)
        elif part:
            # 在普通文本中处理 <sup>...</sup> 上标标签
            _render_sup(paragraph, part)


def _render_sup(paragraph, text):
    """将 <sup>...</sup> 渲染为上标，其余为普通文本"""
    segments = re.split(r"(<sup>.*?</sup>)", text, flags=re.DOTALL)
    for seg in segments:
        if seg.startswith("<sup>") and seg.endswith("</sup>"):
            content = seg[5:-6]  # 去掉标签
            run = paragraph.add_run(content)
            _fmt_run(run, FONT_SONG, size=9)  # 上标字号缩小
            run.font.superscript = True
        elif seg:
            run = paragraph.add_run(seg)
            _fmt_run(run, FONT_SONG, size=12)


# ════════════════════════════════════════════════════════════════════════════
#  段落/标题添加
# ════════════════════════════════════════════════════════════════════════════

def add_paper_title(doc, text):
    """\u8bba\u6587\u5927\u6807\u9898\uff1a\u9ed1\u4f53 16pt \u5c45\u4e2d\u52a0\u7c97"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _fmt_run(run, FONT_HEI, size=16, bold=True)
    _spacing(p, before=24, after=12, lines=1.5)
    return p


def add_abstract_heading(doc):
    """\u6458\u8981\u5c0f\u6807\u9898\uff1a\u9ed1\u4f53 14pt \u5c45\u4e2d"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u6458\u8981")
    _fmt_run(run, FONT_HEI, size=14)
    _spacing(p, before=6, after=4, lines=1.5)
    return p


def add_keywords_line(doc, text):
    """\u5173\u952e\u8bcd\u884c\uff1a\u9ed1\u4f53\u6807\u7b7e + \u5b8b\u4f53\u5185\u5bb9"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _indent(p, first_cm=0.847)
    _spacing(p, before=4, after=0, lines=1.5)
    label = p.add_run("\u5173\u952e\u8bcd\uff1a")
    _fmt_run(label, FONT_HEI, size=12, bold=True)
    content = p.add_run(text.strip())
    _fmt_run(content, FONT_SONG, size=12)
    return p


def parse_abstract(doc, text):
    """
    \u89e3\u6790 00_\u6458\u8981.md\uff1a
      #  -> \u8bba\u6587\u5927\u6807\u9898
      ## \u6458\u8981 -> \u6458\u8981\u6807\u9898 + \u6b63\u6587
      ## \u5173\u952e\u8bcd -> \u5173\u952e\u8bcd\u884c
    """
    mode = None
    abstract_lines = []
    for line in text.split("\n"):
        s = line.strip()
        if not s:
            continue
        if s.startswith("# ") and not s.startswith("## "):
            add_paper_title(doc, s[2:].strip())
            continue
        if s == "## \u6458\u8981":
            add_abstract_heading(doc)
            mode = "abstract"
            continue
        if s == "## \u5173\u952e\u8bcd":
            if abstract_lines:
                add_body(doc, " ".join(abstract_lines))
                abstract_lines = []
            mode = "keywords"
            continue
        if mode == "abstract":
            abstract_lines.append(s)
        elif mode == "keywords":
            add_keywords_line(doc, s)
    if abstract_lines:
        add_body(doc, " ".join(abstract_lines))


def add_h1(doc, text, first=False):
    """一级标题：黑体 15pt 居中，非首章前加分页"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not first:
        _page_break_before(p)
    run = p.add_run(text)
    _fmt_run(run, FONT_HEI, size=15, bold=True)
    _spacing(p, before=12, after=6, lines=1.5)
    return p


def add_h2(doc, text):
    """二级标题：楷体 14pt 左对齐，首行缩进 0.988cm"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _fmt_run(run, FONT_KAI, size=14)
    _spacing(p, before=6, after=3, lines=1.5)
    _indent(p, first_cm=0.988)
    return p


def add_h3(doc, text):
    """三级标题：黑体 12pt 加粗左对齐，首行缩进 0.85cm"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _fmt_run(run, FONT_HEI, size=12, bold=True)
    _spacing(p, before=3, after=0, lines=1.5)
    _indent(p, first_cm=0.85)
    return p


def add_body(doc, text, centered=False):
    """正文段落：宋体 12pt，首行缩进 0.847cm，1.5 倍行距"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    if not centered:
        _indent(p, first_cm=0.847)
    _spacing(p, before=0, after=0, lines=1.5)
    _render_inline(p, text)
    return p


def add_caption(doc, text):
    """图/表标题：宋体 10pt 居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=3, after=3, lines=1.5)
    run = p.add_run(text)
    _fmt_run(run, FONT_SONG, size=10)
    return p


def add_math_block(doc, latex_lines):
    """
    数学块：去除 LaTeX 环境标签，\\\\→换行，居中斜体显示。
    """
    raw = "\n".join(latex_lines)
    raw = re.sub(r"\\begin\{[^}]+\}", "", raw)
    raw = re.sub(r"\\end\{[^}]+\}",   "", raw)
    raw = raw.replace("\\\\", "\n")
    for cl in [l.strip() for l in raw.split("\n") if l.strip()]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, before=2, after=2, lines=1.5)
        run = p.add_run(cl)
        _fmt_run(run, FONT_SONG, size=11, italic=True)


# ════════════════════════════════════════════════════════════════════════════
#  三线表
# ════════════════════════════════════════════════════════════════════════════

def add_three_line_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)

    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri]
        for ci in range(min(len(row_data), ncols)):
            cell = row.cells[ci]
            cell.text = row_data[ci].strip()
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    _fmt_run(run, FONT_SONG, size=10, bold=(ri == 0))

    _apply_three_line(tbl)


def _apply_three_line(table):
    """全表：上框线 1.5pt，下框线 1.5pt，表头行下框线 1.5pt，其余无边框"""
    tbl_el = table._tbl
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl_el.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)

    bds = OxmlElement("w:tblBorders")

    def _line(name, sz="12"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        return el

    def _none(name):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        return el

    bds.append(_line("top"))
    bds.append(_none("left"))
    bds.append(_none("right"))
    bds.append(_line("bottom"))
    bds.append(_none("insideV"))
    bds.append(_none("insideH"))
    tblPr.append(bds)

    # 表头行：每个单元格底部加 1.5pt 线
    if table.rows:
        for cell in table.rows[0].cells:
            tc   = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
            tcBds = tcPr.find(qn("w:tcBorders"))
            if tcBds is None:
                tcBds = OxmlElement("w:tcBorders"); tcPr.append(tcBds)
            b = OxmlElement("w:bottom")
            b.set(qn("w:val"),   "single")
            b.set(qn("w:sz"),    "12")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
            tcBds.append(b)


# ════════════════════════════════════════════════════════════════════════════
#  Markdown 解析
# ════════════════════════════════════════════════════════════════════════════

def parse_md(doc, text, first_h1):
    """
    逐行解析 Markdown，写入 doc。
    返回更新后的 first_h1（bool），跨文件维持"第一个 H1 不分页"状态。
    """
    lines = text.split("\n")
    n = len(lines)
    i = 0
    in_math   = False
    math_buf  = []
    table_buf = []

    def flush_table():
        if table_buf:
            add_three_line_table(doc, table_buf)
            table_buf.clear()

    while i < n:
        line     = lines[i]
        stripped = line.strip()

        # ── 数学块 \[ ... \] ──────────────────────────────────────────────
        if stripped == "\\[":
            flush_table()
            in_math  = True
            math_buf = []
            i += 1
            continue

        if in_math:
            if stripped == "\\]":
                in_math = False
                add_math_block(doc, math_buf)
                math_buf = []
            else:
                math_buf.append(line)
            i += 1
            continue

        # ── 表格行 ────────────────────────────────────────────────────────
        if stripped.startswith("|"):
            # 分隔行 |---|---| 跳过
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_buf.append(cells)
            i += 1
            continue

        # 非表格行时，先刷新缓存的表格
        flush_table()

        # ── 空行 ──────────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── 标题 ──────────────────────────────────────────────────────────
        if stripped.startswith("### "):
            add_h3(doc, stripped[4:].strip())
            i += 1
            continue

        if stripped.startswith("## "):
            add_h2(doc, stripped[3:].strip())
            i += 1
            continue

        if stripped.startswith("# "):
            add_h1(doc, stripped[2:].strip(), first=first_h1)
            first_h1 = False
            i += 1
            continue

        # ── 图/表标题 <center>...</center> ──────────────────────────────
        if stripped.startswith("<center>"):
            caption_text = re.sub(r"<[^>]+>", "", stripped).strip()
            if caption_text:
                add_caption(doc, caption_text)
            i += 1
            continue

        # ── 普通段落（合并连续文本行为一段）──────────────────────────────
        para_lines = [stripped]
        while i + 1 < n:
            nxt = lines[i + 1].strip()
            if (not nxt
                    or nxt.startswith("#")
                    or nxt.startswith("|")
                    or nxt.startswith("<")
                    or nxt == "\\["):
                break
            i += 1
            para_lines.append(lines[i].strip())
        add_body(doc, " ".join(para_lines))
        i += 1

    flush_table()
    return first_h1


# ════════════════════════════════════════════════════════════════════════════
#  页面设置
# ════════════════════════════════════════════════════════════════════════════

def setup_page(doc):
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin   = Cm(3.175)
    sec.right_margin  = Cm(3.175)
    sec.header_distance = Cm(1.50)
    sec.footer_distance = Cm(1.75)


# ════════════════════════════════════════════════════════════════════════════
#  单章生成
# ════════════════════════════════════════════════════════════════════════════

def build_single_chapter(chapter_file):
    """
    生成单个章节的 DOCX 文件
    用法：python build_docx.py 01_绪论.md
    """
    fpath = WRITING_DIR / chapter_file
    if not fpath.exists():
        print(f"[错误] 文件不存在：{fpath}")
        sys.exit(1)

    # 去掉 .md 后缀，拼出输出文件名
    out_name = fpath.stem + ".docx"
    out_path = WRITING_DIR / out_name

    doc = Document()
    setup_page(doc)
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    print(f"处理 {chapter_file} ...")
    content = fpath.read_text(encoding="utf-8")
    parse_md(doc, content, first_h1=True)

    doc.save(str(out_path))
    print(f"\n[OK] 已生成：{out_path}")


# ════════════════════════════════════════════════════════════════════════════
#  主函数
# ════════════════════════════════════════════════════════════════════════════

def main():
    # 如果传入了命令行参数（如 01_绪论.md），只生成这一章
    if len(sys.argv) > 1:
        build_single_chapter(sys.argv[1])
        return

    # 否则生成全文
    doc = Document()
    setup_page(doc)

    # 删除 Document() 自动生成的默认空段落
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    first_h1 = True
    # ── 标题页（摘要）────────────────────────────────────────
    abs_path = WRITING_DIR / ABSTRACT_FILE
    if abs_path.exists():
        print(f"处理 {ABSTRACT_FILE} ...")
        parse_abstract(doc, abs_path.read_text(encoding="utf-8"))
    else:
        print(f"[跳过] {ABSTRACT_FILE} 不存在")

    for fname in CHAPTER_FILES:
        fpath = WRITING_DIR / fname
        if not fpath.exists():
            print(f"[跳过] {fname} 不存在")
            continue
        print(f"处理 {fname} ...")
        content  = fpath.read_text(encoding="utf-8")
        first_h1 = parse_md(doc, content, first_h1)

    # ── 参考文献页 ────────────────────────────────────────────────────────
    ref_h = doc.add_paragraph()
    ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _page_break_before(ref_h)
    _spacing(ref_h, before=12, after=6, lines=1.5)
    run = ref_h.add_run("参考文献")
    _fmt_run(run, FONT_HEI, size=15, bold=True)

    ref_p = doc.add_paragraph()
    _spacing(ref_p, lines=1.15)
    run2 = ref_p.add_run("[需补充参考文献]")
    _fmt_run(run2, FONT_SONG, size=10)

    doc.save(str(OUTPUT_FILE))
    print(f"\n[OK] 已生成：{OUTPUT_FILE}")


if __name__ == "__main__":
    main()
